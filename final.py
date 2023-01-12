import docx
doc = docx.Document()
doc.add_heading('Contrato', 0)

print('Nombre de comprador:')
comprador = input(str())
print('Nombre de vendedor:')
vendedor = input(str())
print('Valor de objeto:')
valor = input(int())

proemio = f'CONTRATO DE COMPRAVENTA QUE CELEBRAN {comprador} Y {vendedor} CONFORME A LAS SIGUIENTES CLAUSULAS'

deccom = 'Declara el comprador:'
decven = 'Declara el vendedor'

objeto = 'OBJETO: El obejto del contrato es una compraventa.'
precio = f'PRECIO: El precio de venta es: ${valor} pesos.'
pago = 'FORMA DE PAGO: El pago se hará a la cuenta de banco 123456789.'
encabezados = 'ENCABEZADOS: Los encabezados son solo de referencia.'
ley = 'LEY APLICABLE Y JURISDICCION: La ley aplicable es la de Monterrey.'

firmas = f'Firma {comprador} (comprador) y Firma {vendedor} (vendedor)'

print('¿Qué contrato quieres hacer?\n A. Contrato de Compraventa')

tipo = input(str())

if tipo == str('a') or str('A'):
    print('Comencemos.')
    print('¿Que clausulas quieres?')
else:
    print('No tenemos ese contrato actualmente.')
    exit()

print("""
1. Proemio\n
2. Declaraciones del comprador\n
3. Declaraciones del Vendedor\n
4. Objeto\n
5. Precio\n
6. Forma de pago\n 
7. Encabezados\n
8. Ley Aplicable\n
9. Firmas\n
Escribe solo los numeros separados con comas.
""")

sel = []  
sel = [int(item) for item in input("Ingresa las clausulas : ").split()]

options = {1: proemio, 2: deccom, 3: decven, 4: objeto, 5: precio, 6: pago, 7: encabezados, 8: ley, 9: firmas}

for option in sel:
    if option in options:
        print(options[option])
        doc_para = doc.add_paragraph(options[option], style='List Number')
    else:
        exit()

doc.save(f'contrato{comprador}.docx')