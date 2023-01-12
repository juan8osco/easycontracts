import docx
doc = docx.Document()
doc.add_heading('Contrato', 0)

print('Nombre de comprador:')
comprador = input(str())
print('Nombre de vendedor:')
vendedor = input(str())
print('Valor de objeto:')
valor = input(int())

proemio = f'CONTRATO DE COMPRAVENTA QUE CELEBRAN {comprador} Y {vendedor} CONFORME A LAS SIGUIENTES CLAUSULAS'

deccom = 'Declara el comprador:'
decven = 'Declara el vendedor'

objeto = 'OBJETO: El obejto del contrato es una compraventa.'
precio = f'PRECIO: El precio de venta es: ${valor} pesos.'
pago = 'FORMA DE PAGO: El pago se hará a la cuenta de banco 123456789.'
encabezados = 'ENCABEZADOS: Los encabezados son solo de referencia.'
ley = 'LEY APLICABLE Y JURISDICCION: La ley aplicable es la de Monterrey.'

firmas = f'Firma {comprador} (comprador) y Firma {vendedor} (vendedor)'

print('¿Qué contrato quieres hacer?\n A. Contrato de Compraventa')

tipo = input(str())

if tipo == str('a') or str('A'):
    print('Comencemos.')
    print('¿Que clausulas quieres?')
else:
    print('No tenemos ese contrato actualmente.')
    exit()

print("""
1. Proemio\n
2. Declaraciones del comprador\n
3. Declaraciones del Vendedor\n
4. Objeto\n
5. Precio\n
6. Forma de pago\n 
7. Encabezados\n
8. Ley Aplicable\n
9. Firmas\n
Escribe solo los numeros separados con comas.
""")

sel = []  
sel = [int(item) for item in input("Ingresa las clausulas : ").split()]

if 1 in sel:
    print(proemio)
    doc_para = doc.add_paragraph(proemio)
if 2 in sel:
    print(deccom)
    doc_para = doc.add_paragraph(deccom)
if 3 in sel:
    print(decven)
    doc_para = doc.add_paragraph(decven)
if 4 in sel:
    print(objeto)
    doc_para = doc.add_paragraph(objeto)
if 5 in sel:
    print(precio)
    doc_para = doc.add_paragraph(precio)
if 6 in sel:
    print(pago)
    doc_para = doc.add_paragraph(pago)
if 7 in sel:
    print(encabezados)
    doc_para = doc.add_paragraph(encabezados)
if 8 in sel:
    print(ley)
    doc_para = doc.add_paragraph(ley)
if 9 in sel:
    print(firmas)
    doc_para = doc.add_paragraph(firmas)
else:
    exit()

doc.save(f'contrato{comprador}.docx')